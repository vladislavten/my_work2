from docx import Document
from docx.shared import Pt  # Добавляем импорт для установки размера шрифта

def replace_text_in_docx(template_path, output_path, replacements):
    doc = Document(template_path)

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
                if key == "vacancy":
                    for run in para.runs:
                        if value in run.text:
                            run.font.name = "Arial"
                            run.font.size = Pt(13)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                        if key == "vacancy":
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if value in run.text:
                                        run.font.name = "Arial"
                                        run.font.size = Pt(13)

    doc.save(output_path)

def main():
    template_path = "Generated_CV.docx"  # Исходный шаблон
    output_path = "Updated_CV.docx"  # Новый файл

    replacements = {
        "Ivan Ivanov": "John Doe",
        "name": "Джон Доууууусон",
        "vacancy": "IT ITITIT",
        "date": "15.06.1990000",
        "Atyrau. Kazakhstan": "New York, USA",
        "500000 tenge on hand": "$8000 per month",
        "Ainur Baituganova, +7 7122 763272 ext.233": "HR Contact, +1 234 567 890",
        "Ainur.Baituganova@fircaspian.com": "hr@example.com",
        "Kazakh National University, Almaty, Kazakhstan": "MIT, Cambridge, USA",
        "Bachelor's Degree in Computer Science": "Master's in Artificial Intelligence",
        "AWS Certified Solutions Architect / 2022": "Google Cloud Certified / 2023",
        "Python, Docker, Kubernetes, AWS, Terraform": "Python, JavaScript, React, Node.js, AWS",
        "Kazakh – Fluent": "Spanish – Fluent",
        "Russian – Native": "English – Native",
        "English – Advanced": "German – Intermediate",
        "Tech Solutions LLC, Kazakhstan": "Big Tech Corp, USA",
        "Senior Backend Engineer": "Software Architect",
        "m.smith@techsolutions.com, +7701XXXXXXX": "contact@bigtech.com, +1 987 654 3210"
    }

    replace_text_in_docx(template_path, output_path, replacements)
    print(f"Updated resume saved as {output_path}")

if __name__ == "__main__":
    main()
